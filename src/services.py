import asyncio
import logging
import os
import tempfile
import subprocess
import io
import yt_dlp
import httpx
import uuid
import json
import requests
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from PIL import Image, ImageDraw, ImageFont

from .config import (
    ASSEMBLYAI_BASE_URL, HEADERS, API_TIMEOUT, FFMPEG_DIR,
    SEGMENT_DURATION, OPENROUTER_API_KEY, FONT_PATH,
    YOOMONEY_WALLET, SUBSCRIPTION_AMOUNT
)

logger = logging.getLogger(__name__)


# =============================
#     YooMoney Payment
# =============================
async def create_yoomoney_payment(user_id: int, amount: int, description: str) -> str:
    """Создает ссылку на оплату YooMoney."""
    payment_label = f"sub_{user_id}_{uuid.uuid4()}"
    quickpay_url = "https://yoomoney.ru/quickpay/confirm.xml"
    params = {
        "receiver": YOOMONEY_WALLET,
        "quickpay-form": "shop",
        "targets": description,
        "paymentType": "SB",
        "sum": amount,
        "label": payment_label,
    }
    
    async with httpx.AsyncClient() as client:
        try:
            # The POST request is for validation. A 302 redirect is expected and not an error.
            await client.post(quickpay_url, data=params)
            
            # YooMoney QuickPay form doesn't return a JSON with a URL,
            # it redirects. We build the URL for the user to follow.
            # The above POST is more for validation/logging on YooMoney's side.
            # The actual payment link is constructed with GET parameters.
            
            from urllib.parse import urlencode
            encoded_params = urlencode(params)
            payment_url = f"https://yoomoney.ru/quickpay/confirm.xml?{encoded_params}"
            
            logger.info(f"Создана ссылка на оплату для user_id {user_id}: {payment_label}")
            return payment_url, payment_label

        except httpx.RequestError as e:
            logger.error(f"Ошибка при создании платежа YooMoney для user_id {user_id}: {e}")
            return None, None


# =============================
#     Регистрация шрифта PDF
# =============================
try:
    # Try to register DejaVu font first
    if os.path.exists(FONT_PATH):
        pdfmetrics.registerFont(TTFont("DejaVu", FONT_PATH))
        logger.info(f"Successfully registered DejaVu font: {FONT_PATH}")
    else:
        logger.warning(f"DejaVu font not found: {FONT_PATH}")

    # Register Noto Sans font for excellent Cyrillic/Unicode support
    noto_path = os.path.join(os.path.dirname(FONT_PATH), "NotoSans-Regular.ttf")
    if os.path.exists(noto_path):
        pdfmetrics.registerFont(TTFont("NotoSans", noto_path))
        logger.info(f"Successfully registered NotoSans font: {noto_path}")
    else:
        logger.warning(f"NotoSans font not found: {noto_path}")

    # Register Arial font as fallback
    arial_path = os.path.join(os.path.dirname(FONT_PATH), "arial.ttf")
    if os.path.exists(arial_path):
        pdfmetrics.registerFont(TTFont("Arial", arial_path))
        logger.info(f"Successfully registered Arial font: {arial_path}")
    else:
        logger.warning(f"Arial font not found: {arial_path}, using default fonts")
except Exception as e:
    logger.error(f"Failed to register custom fonts: {e}, using default fonts")

# ---------- Сохранение в разные форматы ----------

def _register_pdf_font_if_needed():
    try:
        # Register NotoSans if not already registered (best Unicode support)
        noto_path = os.path.join(os.path.dirname(FONT_PATH), "NotoSans-Regular.ttf")
        if 'NotoSans' not in pdfmetrics.getRegisteredFontNames() and os.path.exists(noto_path):
            pdfmetrics.registerFont(TTFont("NotoSans", noto_path))

        # Register Arial if not already registered
        arial_path = os.path.join(os.path.dirname(FONT_PATH), "arial.ttf")
        if 'Arial' not in pdfmetrics.getRegisteredFontNames() and os.path.exists(arial_path):
            pdfmetrics.registerFont(TTFont("Arial", arial_path))

        # Register DejaVu if not already registered
        if 'DejaVu' not in pdfmetrics.getRegisteredFontNames() and os.path.exists(FONT_PATH):
            pdfmetrics.registerFont(TTFont("DejaVu", FONT_PATH))
    except Exception:
        pass

def save_text_to_pdf(text: str, output_path: str):
    _register_pdf_font_if_needed()
    doc = SimpleDocTemplate(output_path, pagesize=A4,
                            rightMargin=50, leftMargin=50,
                            topMargin=50, bottomMargin=50)
    styles = getSampleStyleSheet()

    # Ensure proper encoding for Cyrillic text
    from reportlab.pdfbase import pdfdoc
    pdfdoc.ENCODING = 'UTF-8'

    style = styles['Normal']
    # Use NotoSans if available (best Unicode support), otherwise Arial, then DejaVu, then defaults
    available_fonts = pdfmetrics.getRegisteredFontNames()
    if 'NotoSans' in available_fonts:
        style.fontName = 'NotoSans'
    elif 'Arial' in available_fonts:
        style.fontName = 'Arial'
    elif 'DejaVu' in available_fonts:
        style.fontName = 'DejaVu'
    else:
        # Try fonts that typically support Cyrillic/Unicode
        preferred_fonts = ['Times-Roman', 'Courier', 'Times-Bold', 'Courier-Bold']
        style.fontName = next((f for f in preferred_fonts if f in available_fonts), 'Times-Roman')
    style.fontSize = 12
    style.leading = 15

    # Ensure text is properly handled as Unicode
    if isinstance(text, bytes):
        text = text.decode('utf-8', errors='replace')
    elif not isinstance(text, str):
        text = str(text)

    # Clean and normalize the text
    import unicodedata
    text = unicodedata.normalize('NFC', text)

    paragraphs = [Paragraph(p.replace('\n', '<br />'), style) for p in text.split('\n\n') if p.strip()]
    elems = []
    for p in paragraphs:
        elems.append(p)
        elems.append(Spacer(1, 12))
    doc.build(elems)


def save_text_to_txt(text: str, output_path: str):
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(text)


def save_text_to_md(text: str, output_path: str):
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(text)


def save_text_to_docx(text: str, output_path: str):
    try:
        from docx import Document
        doc = Document()
        for par in text.split("\n\n"):
            for line in par.split("\n"):
                doc.add_paragraph(line)
            doc.add_paragraph("")
        doc.save(output_path)
    except Exception as e:
        logger.warning(f"Не удалось сохранить DOCX ({e}), сохраняю как TXT")
        save_text_to_txt(text, output_path)


# ---------- Аудио-обработка / API ----------
class AudioProcessor:
    @staticmethod
    def split_audio(input_path: str, segment_time: int = SEGMENT_DURATION) -> list[str]:
        output_dir = tempfile.mkdtemp(prefix="fragments_")
        output_pattern = os.path.join(output_dir, "fragment_%03d.mp3")
        ffmpeg_path = os.path.join(FFMPEG_DIR, "ffmpeg.exe")
        command = [
            ffmpeg_path,
            "-i", input_path,
            "-f", "segment",
            "-segment_time", str(segment_time),
            "-c", "copy",
            output_pattern
        ]
        try:
            subprocess.run(command, check=True, capture_output=True, text=True)
            return sorted([
                os.path.join(output_dir, f)
                for f in os.listdir(output_dir)
                if f.startswith("fragment_") and f.endswith(".mp3")
            ])
        except subprocess.CalledProcessError as e:
            logger.error(f"FFmpeg error: {e.stderr}")
            raise RuntimeError("Ошибка при разделении аудио") from e

    @staticmethod
    def cleanup(files: list[str]):
        for path in files:
            try:
                if os.path.isfile(path):
                    os.remove(path)
                elif os.path.isdir(path):
                    for f in os.listdir(path):
                        os.remove(os.path.join(path, f))
                    os.rmdir(path)
            except Exception as e:
                logger.warning(f"Ошибка удаления {path}: {e}")


async def upload_to_assemblyai(file_path: str, retries: int = 3) -> str:
    for attempt in range(retries):
        try:
            async with httpx.AsyncClient() as client:
                with open(file_path, "rb") as f:
                    response = await client.post(
                        f"{ASSEMBLYAI_BASE_URL}/upload",
                        headers=HEADERS,
                        files={"file": f},
                        timeout=API_TIMEOUT
                    )
                response.raise_for_status()
                return response.json()["upload_url"]
        except Exception as e:
            logger.warning(f"Попытка {attempt + 1}/{retries} загрузки файла не удалась: {str(e)}")
            if attempt == retries - 1:
                raise RuntimeError("Не удалось загрузить файл на сервер AssemblyAI") from e
            await asyncio.sleep(2 ** attempt)


async def transcribe_with_assemblyai(audio_url: str, retries: int = 3) -> dict:
    headers = {
        "authorization": HEADERS['authorization'],
        "content-type": "application/json"
    }
    payload = {
        "audio_url": audio_url,
        "speaker_labels": True,
        "punctuate": True,
        "format_text": True,
        "language_code": "ru",  # Explicitly set Russian language
        "language_detection": False  # Disable auto-detection since we specify Russian
    }
    for attempt in range(retries):
        try:
            async with httpx.AsyncClient() as client:
                resp = await client.post(
                    "https://api.assemblyai.com/v2/transcript",
                    headers=headers, json=payload
                )
                resp.raise_for_status()
                transcript_id = resp.json()["id"]
                while True:
                    status = await client.get(
                        f"https://api.assemblyai.com/v2/transcript/{transcript_id}",
                        headers=headers
                    )
                    result = status.json()
                    if result["status"] == "completed":
                        return result
                    elif result["status"] == "error":
                        raise Exception(result["error"])
                    await asyncio.sleep(3)
        except Exception as e:
            logger.warning(f"Попытка {attempt + 1}/{retries} транскрипции не удалась: {str(e)}")
            if attempt == retries - 1:
                raise
            await asyncio.sleep(2 ** attempt)


async def download_youtube_audio(url: str, progress_callback: callable = None) -> str:
    loop = asyncio.get_running_loop()
    progress_queue = asyncio.Queue()

    def progress_hook(data):
        if data['status'] == 'downloading' and progress_callback:
            try:
                percent_str = data.get('_percent_str', '0%')
                percent_value = float(percent_str.strip().replace('%', ''))
                loop.call_soon_threadsafe(progress_queue.put_nowait, percent_value)
            except:
                pass

    def sync_download():
        temp_dir = tempfile.gettempdir()
        unique_id = str(uuid.uuid4())
        outtmpl = os.path.join(temp_dir, f"{unique_id}")
        ydl_opts = {
            "format": "bestaudio/best",
            "outtmpl": outtmpl,
            "ffmpeg_location": FFMPEG_DIR,
            "progress_hooks": [progress_hook],
            "postprocessors": [{
                "key": "FFmpegExtractAudio",
                "preferredcodec": "mp3",
            }],
            "quiet": True,
            "no_warnings": False,
            "extract_flat": False,
            "http_headers": {
                "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36",
            },
        }
        try:
            with yt_dlp.YoutubeDL(ydl_opts) as ydl:
                ydl.download([url])
                expected_filename = f"{outtmpl}.mp3"
                if os.path.exists(expected_filename):
                    return expected_filename
                else:
                    for file in os.listdir(temp_dir):
                        if file.startswith(os.path.basename(unique_id)):
                            return os.path.join(temp_dir, file)
                    raise FileNotFoundError(f"Скачанный аудиофайл не найден: {expected_filename}")
        except Exception as e:
            logger.error(f"Ошибка скачивания YouTube: {str(e)}")
            raise RuntimeError(f"Ошибка скачивания видео: {str(e)}") from e

    async def process_progress():
        while True:
            try:
                data = await asyncio.wait_for(progress_queue.get(), timeout=1.0)
                if progress_callback:
                    await progress_callback(data)
            except asyncio.TimeoutError:
                if download_task.done():
                    break
            except Exception as e:
                logger.warning(f"Ошибка обработки прогресса: {str(e)}")
                break

    download_task = loop.run_in_executor(None, sync_download)
    progress_task = asyncio.create_task(process_progress())
    try:
        result = await download_task
        return result
    finally:
        progress_task.cancel()
        try:
            await progress_task
        except asyncio.CancelledError:
            pass


def format_results_with_speakers(segments: list[dict]) -> str:
    return "\n\n".join(f"Спикер {seg['speaker']}:\n{seg['text']}" for seg in segments)


def format_results_plain(segments: list[dict]) -> str:
    return "\n\n".join(seg["text"] for seg in segments)


def generate_summary_timecodes(segments: list[dict]) -> str:
    url = "https://openrouter.ai/api/v1/chat/completions"
    headers = {
        "Authorization": f"Bearer {OPENROUTER_API_KEY}",
        "Content-Type": "application/json"
    }
    full_text_with_timestamps = ""
    for i, seg in enumerate(segments):
        start_minute = i * SEGMENT_DURATION // 60
        start_second = i * SEGMENT_DURATION % 60
        start_code = f"{start_minute:02}:{start_second:02}"
        full_text_with_timestamps += f"[{start_code}] {seg['text']}\n\n"
    prompt = f"""
Проанализируй полную расшифровку аудио с тайм-кодами и создай структурированное оглавление.
Текст с тайм-кодами:
{full_text_with_timestamps}
Инструкции:
1. Выдели ОСНОВНЫЕ смысловые блоки и темы
2. Группируй несколько последовательных сегментов в один логический блок
3. Для каждого блока укажи время начала
4. Дай емкое описание содержания блока
5. Сохраняй хронологический порядок
Формат ответа:
Тайм-коды
MM:SS - [Основная тема/событие]
[Дополнительные детали]
MM:SS - [Следующая основная тема]
...
"""
    data = {
        "model": "z-ai/glm-4.5-air:free",
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.2
    }
    try:
        response = requests.post(url, headers=headers, data=json.dumps(data), timeout=60)
        response.raise_for_status()
        return response.json()['choices'][0]['message']['content'].strip()
    except Exception:
        fallback_result = "Тайм-коды\n\n"
        for i, seg in enumerate(segments):
            start_minute = i * SEGMENT_DURATION // 60
            start_second = i * SEGMENT_DURATION % 60
            start_code = f"{start_minute:02}:{start_second:02}"
            fallback_result += f"{start_code} - {seg['text'][:50]}...\n"
        return fallback_result


async def convert_to_mp3(input_path: str) -> str:
    output_path = tempfile.NamedTemporaryFile(delete=False, suffix=".mp3").name
    ffmpeg_path = os.path.join(FFMPEG_DIR, "ffmpeg.exe")
    command = [
        ffmpeg_path,
        "-i", input_path,
        "-acodec", "libmp3lame",
        "-q:a", "2",
        "-y",
        output_path
    ]
    try:
        process = await asyncio.create_subprocess_exec(
            *command,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE
        )
        stdout, stderr = await process.communicate()
        if process.returncode != 0:
            logger.error(f"Ошибка конвертации: {stderr.decode()}")
            raise RuntimeError(f"Ошибка конвертации файла: {stderr.decode()}")
        if not os.path.exists(output_path) or os.path.getsize(output_path) == 0:
            logger.error(f"Выходной файл {output_path} не создан или пуст")
            raise RuntimeError("Конвертация не удалась: выходной файл не создан")
        logger.info(f"Конвертация успешна: {output_path}")
        return output_path
    except Exception as e:
        logger.error(f"Ошибка конвертации {input_path}: {str(e)}")
        raise RuntimeError(f"Ошибка конвертации: {str(e)}") from e


async def process_audio_file(file_path: str, user_id: int, progress_callback=None) -> list[dict]:
    try:
        logger.info(f"Обработка аудиофайла: {file_path}")
        if progress_callback:
            await progress_callback(0.01, "Загружаю файл для обработки...")
        audio_url = await upload_to_assemblyai(file_path)
        if progress_callback:
            await progress_callback(0.30, "Запускаю транскрибацию...")
        result = await transcribe_with_assemblyai(audio_url)
        if progress_callback:
            await progress_callback(0.90, "Формирую результаты...")

        # Log the raw result for debugging
        logger.info(f"AssemblyAI raw result keys: {list(result.keys())}")
        if "language_code" in result:
            logger.info(f"Detected language: {result['language_code']}")

        segments = []
        if "utterances" in result and result["utterances"]:
            for utt in result["utterances"]:
                text = (utt.get("text") or "").strip()
                # Ensure proper UTF-8 encoding
                if isinstance(text, str):
                    text = text.encode('utf-8').decode('utf-8')
                segments.append({
                    "speaker": utt.get("speaker", "?"),
                    "text": text
                })
                logger.debug(f"Segment: speaker={utt.get('speaker')}, text_length={len(text)}")
        elif "text" in result:
            text = (result["text"] or "").strip()
            # Ensure proper UTF-8 encoding
            if isinstance(text, str):
                text = text.encode('utf-8').decode('utf-8')
            segments.append({"speaker": "?", "text": text})
            logger.debug(f"Full text length: {len(text)}")

        if progress_callback:
            await progress_callback(1.0, "Обработка завершена!")
        logger.info(f"Транскрибация завершена, найдено {len(segments)} сегментов")
        return segments
    except Exception as e:
        logger.error(f"Ошибка в process_audio_file: {str(e)}")
        raise


THUMBNAIL_CACHE = {}

def create_custom_thumbnail(thumbnail_path: str = None):
    cache_key = thumbnail_path or "default"
    if cache_key in THUMBNAIL_CACHE:
        thumbnail_bytes = io.BytesIO(THUMBNAIL_CACHE[cache_key])
        thumbnail_bytes.seek(0)
        return thumbnail_bytes
    try:
        if thumbnail_path and os.path.exists(thumbnail_path):
            with Image.open(thumbnail_path) as img:
                if img.mode == 'RGBA':
                    background = Image.new('RGB', img.size, (255, 255, 255))
                    background.paste(img, mask=img.split()[3])
                    img = background
                target_size = (320, 320)
                img.thumbnail(target_size, Image.LANCZOS)
                square_img = Image.new('RGB', target_size, (255, 255, 255))
                x_offset = (target_size[0] - img.width) // 2
                y_offset = (target_size[1] - img.height) // 2
                square_img.paste(img, (x_offset, y_offset))
                thumbnail_bytes = io.BytesIO()
                square_img.save(thumbnail_bytes, format='JPEG', quality=95, optimize=True)
                thumbnail_bytes.seek(0)
                THUMBNAIL_CACHE[cache_key] = thumbnail_bytes.getvalue()
                thumbnail_bytes.seek(0)
                return thumbnail_bytes
        else:
            target_size = (320, 320)
            img = Image.new('RGB', target_size, color=(230, 50, 50))
            draw = ImageDraw.Draw(img)
            margin = 10
            draw.rectangle([margin, margin, target_size[0] - margin, target_size[1] - margin],
                           outline=(255, 255, 255), width=4)
            try:
                font = ImageFont.truetype("arial.ttf", 80)
            except:
                font = ImageFont.load_default()
            text = "PDF"
            bbox = draw.textbbox((0, 0), text, font=font)
            text_width = bbox[2] - bbox[0]
            text_height = bbox[3] - bbox[1]
            x = (target_size[0] - text_width) // 2
            y = (target_size[1] - text_height) // 2
            draw.text((x, y), text, fill=(255, 255, 255), font=font)
            thumbnail_bytes = io.BytesIO()
            img.save(thumbnail_bytes, format='JPEG', quality=95, optimize=True)
            thumbnail_bytes.seek(0)
            THUMBNAIL_CACHE[cache_key] = thumbnail_bytes.getvalue()
            thumbnail_bytes.seek(0)
            return thumbnail_bytes
    except Exception as e:
        logger.error(f"Ошибка создания thumbnail: {e}")
        return None
